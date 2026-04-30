"""
generators/docx_generator.py
Converts a ParsedDoc into a beautifully formatted, fully editable .docx.

Features:
  • Proper Heading 1–4 styles (Word Document Map / Navigation Pane)
  • Native bullet + numbered list styles
  • Inline bold, italic, equation (Consolas + blue colour) formatting
  • Shaded blockquotes with left border
  • Diagram callout boxes (amber shade)
  • Full Markdown table rendering as Word tables
  • Metadata header line + footer rule
"""

import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agent.parser import ParsedDoc, Block, BT, Seg
from config import APP_NAME, APP_VERSION


# colour palette
_C1 = RGBColor(0x1F, 0x49, 0x7D)   # dark navy
_C2 = RGBColor(0x2E, 0x74, 0xB5)   # medium blue
_C3 = RGBColor(0x5A, 0x96, 0xC8)   # light blue
_CM = RGBColor(0x60, 0x60, 0x60)   # muted grey
_CE = RGBColor(0x1A, 0x3A, 0x6E)   # equation ink


class DocxGenerator:

    def generate(self, doc: ParsedDoc, path: str) -> None:
        wd = Document()
        self._page_setup(wd)
        self._heading_styles(wd)
        self._title_block(wd, doc)
        for block in doc.blocks:
            self._block(wd, block)
        self._footer(wd, doc)
        wd.save(path)

    # ── page & base styles ───────────────────────────────────
    def _page_setup(self, wd):
        sec = wd.sections[0]
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin = sec.right_margin  = Cm(2.5)
        sec.top_margin  = sec.bottom_margin = Cm(2.0)

        normal = wd.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

    def _heading_styles(self, wd):
        specs = [
            ("Heading 1", 18, _C1, True),
            ("Heading 2", 14, _C2, True),
            ("Heading 3", 12, _C3, True),
            ("Heading 4", 11, _CM, True),
        ]
        for name, sz, clr, bold in specs:
            try:
                s = wd.styles[name]
            except KeyError:
                s = wd.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            s.font.name  = "Calibri"
            s.font.size  = Pt(sz)
            s.font.bold  = bold
            s.font.color.rgb = clr
            s.paragraph_format.space_before = Pt(10)
            s.paragraph_format.space_after  = Pt(4)

    # ── title block ──────────────────────────────────────────
    def _title_block(self, wd, doc: ParsedDoc):
        p = wd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(doc.title)
        r.font.name  = "Calibri"
        r.font.size  = Pt(22)
        r.font.bold  = True
        r.font.color.rgb = _C1

        sub = wd.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(
            f"Extracted by {APP_NAME} v{APP_VERSION}  ·  "
            f"{datetime.now().strftime('%d %b %Y  %H:%M')}"
        )
        sr.font.size  = Pt(9)
        sr.font.color.rgb = _CM
        sr.italic = True

        self._hrule(wd)

    # ── block dispatcher ─────────────────────────────────────
    def _block(self, wd, block: Block):
        bt = block.btype
        if   bt == BT.H1:       self._heading(wd, block.raw, "Heading 1")
        elif bt == BT.H2:       self._heading(wd, block.raw, "Heading 2")
        elif bt == BT.H3:       self._heading(wd, block.raw, "Heading 3")
        elif bt == BT.H4:       self._heading(wd, block.raw, "Heading 4")
        elif bt == BT.PARA:     self._para(wd, block.segs)
        elif bt == BT.BULLETS:  self._list_block(wd, block, ordered=False)
        elif bt == BT.NUMBERED: self._list_block(wd, block, ordered=True)
        elif bt == BT.EQUATION: self._equation(wd, block.raw)
        elif bt == BT.DIAGRAM:  self._diagram(wd, block.raw)
        elif bt == BT.QUOTE:    self._quote(wd, block.raw, block.segs)
        elif bt == BT.TABLE:    self._table(wd, block.raw)
        elif bt == BT.DIVIDER:  self._hrule(wd)

    # ── paragraph writers ────────────────────────────────────
    def _heading(self, wd, text, style):
        p = wd.add_paragraph(style=style)
        p.add_run(text.lstrip("#").strip())

    def _para(self, wd, segs: list[Seg]):
        if not segs:
            return
        p = wd.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        for seg in segs:
            self._run(p, seg)

    def _list_block(self, wd, block: Block, ordered: bool):
        style = "List Number" if ordered else "List Bullet"
        for idx, child in enumerate(block.children, 1):
            try:
                p = wd.add_paragraph(style=style)
            except KeyError:
                p = wd.add_paragraph()
                p.add_run(f"{idx}. " if ordered else "• ")
            for seg in child.segs:
                self._run(p, seg)

    def _equation(self, wd, eq: str):
        p = wd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        _shade(p, "EFF4FB")
        r = p.add_run(f"  {eq}  ")
        r.font.name  = "Consolas"
        r.font.size  = Pt(11)
        r.font.color.rgb = _CE

    def _diagram(self, wd, desc: str):
        p = wd.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        _shade(p, "FFF8E7")
        lbl = p.add_run("📊 Diagram: ")
        lbl.bold = True
        lbl.font.color.rgb = RGBColor(0xD2, 0x99, 0x22)
        p.add_run(desc)

    def _quote(self, wd, raw: str, segs: list[Seg]):
        p = wd.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1.0)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        _shade(p, "F0F8FF")
        _left_border(p, "2E74B5")
        for seg in (segs or [Seg(raw)]):
            self._run(p, seg)

    def _table(self, wd, raw: str):
        lines = [l.strip() for l in raw.splitlines() if l.strip()]
        rows  = []
        for line in lines:
            if re.match(r'\|[-| :]+\|', line):
                continue
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)
        if not rows:
            return

        ncols = max(len(r) for r in rows)
        rows  = [r + [""] * (ncols - len(r)) for r in rows]
        tbl   = wd.add_table(rows=len(rows), cols=ncols)
        tbl.style = "Table Grid"

        for ri, row in enumerate(rows):
            for ci, cell_text in enumerate(row):
                cell = tbl.cell(ri, ci)
                cell.text = cell_text
                if ri == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.color.rgb = _C1

        wd.add_paragraph()

    # ── inline run ───────────────────────────────────────────
    def _run(self, p, seg: Seg):
        r = p.add_run(seg.text)
        r.font.name = "Calibri"
        if seg.bold:    r.bold   = True
        if seg.italic:  r.italic = True
        if seg.is_eq:
            r.font.name  = "Consolas"
            r.font.size  = Pt(10.5)
            r.font.color.rgb = _CE
        if seg.is_code:
            r.font.name = "Consolas"
            r.font.size = Pt(10)

    # ── footer ───────────────────────────────────────────────
    def _footer(self, wd, doc: ParsedDoc):
        self._hrule(wd)
        p = wd.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            f"Generated by {APP_NAME} v{APP_VERSION}  ·  FAST-NUCES  ·  "
            f"{doc.stats.get('words', 0)} words  ·  "
            f"{doc.stats.get('equations', 0)} equations"
        )
        r.font.size = Pt(8)
        r.font.color.rgb = _CM
        r.italic = True

    # ── XML helpers ───────────────────────────────────────────
    def _hrule(self, wd):
        p   = wd.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr= OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "AAAAAA")
        pBdr.append(bot)
        pPr.append(pBdr)


def _shade(p, hex_color: str):
    try:
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        pPr.append(shd)
    except Exception:
        pass

def _left_border(p, hex_color: str):
    try:
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"),   "single")
        left.set(qn("w:sz"),    "24")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), hex_color)
        pBdr.append(left)
        pPr.append(pBdr)
    except Exception:
        pass
